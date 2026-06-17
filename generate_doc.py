
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=None):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = heading.runs[0]
    if color:
        run.font.color.rgb = RGBColor(*color)
    run.font.bold = True
    return heading

def add_paragraph(doc, text, bold=False, italic=False, size=11, color=None, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def add_table_row(table, cells_data, is_header=False):
    row = table.add_row()
    for i, (text, width) in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = text
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_header else WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.font.size = Pt(10)
                run.font.bold = is_header
                if is_header:
                    run.font.color.rgb = RGBColor(255, 255, 255)
        if is_header:
            set_cell_bg(cell, '1F3864')
    return row

# ─────────────────────────────────────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Default body font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ─── COVER / TITLE ───────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('GAME OPS SYSTEM')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(31, 56, 100)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run('Technical Design & Implementation Document')
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(80, 80, 80)
run2.italic = True

doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Multiplayer Action Game — Weekend Event Operations\n').font.size = Pt(11)
meta.add_run('Approach: Hybrid Rule-Based Python + FastAPI + Interactive Web Dashboard').font.size = Pt(11)
doc.add_paragraph()

doc.add_page_break()

# ─── 1. PROBLEM UNDERSTANDING ────────────────────────────────────────────────
add_heading(doc, '1. Problem Understanding', level=1, color=(31, 56, 100))

add_paragraph(doc,
    'During a weekend multiplayer gaming event, thousands of players submit match results in real time. '
    'The game operations team needs a system that can process incoming telemetry, maintain a fair competitive '
    'leaderboard, detect suspicious or anomalous player behaviour (cheating), group players into balanced '
    'match lobbies, and explain how the system scales as player count grows from 2,000 to 200,000+.',
    size=11)

doc.add_paragraph()
add_paragraph(doc, 'Core Questions Addressed:', bold=True)
add_bullet(doc, 'What data do we collect and how do we clean it?')
add_bullet(doc, 'How do we rank players fairly, even when some scores appear fraudulent?')
add_bullet(doc, 'How do we detect cheaters without training data or ML models?')
add_bullet(doc, 'How do we group players into fair, low-latency match lobbies?')
add_bullet(doc, 'How does the architecture scale for very high player volumes?')

doc.add_paragraph()

# ─── 2. ASSUMPTIONS ──────────────────────────────────────────────────────────
add_heading(doc, '2. Assumptions', level=1, color=(31, 56, 100))

assumptions = [
    ('Data Source',     'CSV file (or POST API call) as the ingestion mechanism. No real-time stream assumed at prototype level.'),
    ('Thresholds',      'Anti-cheat rule thresholds are manually tuned based on sample data analysis; they are configurable in config.'),
    ('Fair Leaderboard','A "fair" leaderboard excludes or visibly separates cheaters. Flagged players still appear but are marked "under_review" — their scores do NOT push any clean player down in rank.'),
    ('Matchmaking',     'Static bucketing is sufficient at prototype scale. A live queue system with wait-time optimization is a production enhancement.'),
    ('No Auth',         'API endpoints have no authentication or rate limiting at prototype level. This would be added in production.'),
    ('Single Region DB','In-memory Pandas DataFrame used as the datastore. Production would use PostgreSQL with Redis cache.'),
]

tbl = doc.add_table(rows=1, cols=2)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tbl.rows[0].cells
hdr[0].text = 'Assumption Area'
hdr[1].text = 'Detail'
for cell in hdr:
    set_cell_bg(cell, '1F3864')
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255,255,255)
            run.font.size = Pt(10)

for area, detail in assumptions:
    row = tbl.add_row()
    row.cells[0].text = area
    row.cells[1].text = detail
    row.cells[0].paragraphs[0].runs[0].font.bold = True
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)
    set_cell_bg(row.cells[0], 'DCE6F1')

doc.add_paragraph()

# ─── 3. DATA STRUCTURE & DATABASE DESIGN ─────────────────────────────────────
add_heading(doc, '3. Data Structure & Database Design', level=1, color=(31, 56, 100))

add_paragraph(doc, '3.1  Raw Input Schema (matches table)', bold=True)
add_paragraph(doc,
    'Each match event contains the following fields. Two derived columns are computed immediately '
    'on ingestion and stored alongside raw data:')

fields = [
    ('player_id',               'VARCHAR',  'Unique player identifier (Primary Key Part 1)'),
    ('match_id',                'VARCHAR',  'Unique match identifier (Primary Key Part 2)'),
    ('region',                  'VARCHAR',  'Player region — India, SEA, Europe, etc.'),
    ('device',                  'VARCHAR',  'Device type — Android, iOS, PC, Console'),
    ('ping',                    'INT',      'Network latency in milliseconds'),
    ('score',                   'FLOAT',    'Final match score'),
    ('kills',                   'INT',      'Number of kills in this match'),
    ('deaths',                  'INT',      'Number of deaths in this match'),
    ('match_duration_seconds',  'INT',      'Match length in seconds (must be > 0)'),
    ('score_per_second',        'FLOAT',    'DERIVED — score ÷ match_duration_seconds'),
    ('kd_ratio',                'FLOAT',    'DERIVED — kills ÷ max(deaths, 1)'),
    ('created_at',              'TIMESTAMP','Ingestion timestamp (auto-set by server)'),
]

tbl2 = doc.add_table(rows=1, cols=3)
tbl2.style = 'Table Grid'
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, txt in enumerate(['Field', 'Type', 'Description']):
    cell = tbl2.rows[0].cells[i]
    cell.text = txt
    set_cell_bg(cell, '2E75B6')
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255,255,255)
            run.font.size = Pt(10)

for fname, ftype, fdesc in fields:
    row = tbl2.add_row()
    row.cells[0].text = fname
    row.cells[1].text = ftype
    row.cells[2].text = fdesc
    is_derived = 'DERIVED' in fdesc
    bg = 'FFF2CC' if is_derived else 'FFFFFF'
    for c in row.cells:
        set_cell_bg(c, bg)
        c.paragraphs[0].runs[0].font.size = Pt(10)
    row.cells[0].paragraphs[0].runs[0].font.bold = True

doc.add_paragraph()
add_paragraph(doc, '3.2  Player Aggregates Schema (players table)', bold=True)
add_paragraph(doc,
    'A summary table is maintained per player, updated after each match batch. '
    'It is used by leaderboard and matchmaking modules directly:')

pagg = [
    ('player_id',           'Unique player (Primary Key)'),
    ('total_score',         'Sum of all match scores'),
    ('total_kills',         'Sum of all kills'),
    ('total_deaths',        'Sum of all deaths'),
    ('matches_played',      'Count of processed matches'),
    ('primary_region',      'Mode of regions played in'),
    ('skill_score',         'Computed weighted skill metric'),
    ('cheat_status',        '"clean" or "under_review"'),
    ('cheat_severity',      '"None", "Low", "Medium", or "High"'),
    ('flag_reasons',        'Semicolon-delimited list of triggered rules'),
]

tbl3 = doc.add_table(rows=1, cols=2)
tbl3.style = 'Table Grid'
for i, txt in enumerate(['Field', 'Description']):
    cell = tbl3.rows[0].cells[i]
    cell.text = txt
    set_cell_bg(cell, '2E75B6')
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255,255,255)
            run.font.size = Pt(10)

for fname, fdesc in pagg:
    row = tbl3.add_row()
    row.cells[0].text = fname
    row.cells[1].text = fdesc
    for c in row.cells:
        c.paragraphs[0].runs[0].font.size = Pt(10)
    row.cells[0].paragraphs[0].runs[0].font.bold = True

doc.add_paragraph()

# ─── 4. LEADERBOARD LOGIC ────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, '4. Leaderboard Logic', level=1, color=(31, 56, 100))

add_paragraph(doc, '4.1  Aggregation Step', bold=True)
add_paragraph(doc,
    'Before ranking, all match records for each player are collapsed into a single '
    'aggregate row using the following accumulations:')
for t in ['total_score = SUM(score)',
          'total_kills = SUM(kills)',
          'total_deaths = SUM(deaths)',
          'matches_played = COUNT(match_id)',
          'avg_score_per_match = total_score ÷ matches_played',
          'primary_region = MODE(region)  ← most frequently played region']:
    add_bullet(doc, t)

doc.add_paragraph()
add_paragraph(doc, '4.2  Tie-Breaker Priority (applied in order)', bold=True)

rules = [
    ('1st', 'total_score',    'Descending', 'Higher cumulative score → better rank'),
    ('2nd', 'total_deaths',   'Ascending',  'Fewer deaths → better rank (same score)'),
    ('3rd', 'total_kills',    'Descending', 'More kills → better rank (same score + deaths)'),
    ('4th', 'matches_played', 'Descending', 'More matches played → better rank (all else equal)'),
]

tbl4 = doc.add_table(rows=1, cols=4)
tbl4.style = 'Table Grid'
for i, txt in enumerate(['Priority', 'Column', 'Order', 'Rationale']):
    cell = tbl4.rows[0].cells[i]
    cell.text = txt
    set_cell_bg(cell, '375623')
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255,255,255)
            run.font.size = Pt(10)

for pri, col, order, rationale in rules:
    row = tbl4.add_row()
    row.cells[0].text = pri
    row.cells[1].text = col
    row.cells[2].text = order
    row.cells[3].text = rationale
    for c in row.cells:
        c.paragraphs[0].runs[0].font.size = Pt(10)
    row.cells[0].paragraphs[0].runs[0].font.bold = True

doc.add_paragraph()
add_paragraph(doc, '4.3  Fairness — Handling Flagged Players', bold=True)
add_paragraph(doc,
    'Players identified by the anti-cheat module are handled transparently: '
    'they appear at the bottom of the leaderboard table with status = "under_review" '
    'and rank = None. This ensures:')
add_bullet(doc, 'Fraudulent scores do NOT displace any clean player\'s rank number.')
add_bullet(doc, 'Operations teams can still see flagged player statistics for manual review.')
add_bullet(doc, 'The competitive integrity of ranked positions 1–N is fully protected.')

doc.add_paragraph()
add_paragraph(doc, '4.4  Regional Leaderboard', bold=True)
add_paragraph(doc,
    'A separate region-wise leaderboard is generated by grouping clean players by their '
    'primary_region and applying the identical sort/tie-breaker logic within each group. '
    'This allows players to compare their standing within their own region independently '
    'of the global ranking.')

doc.add_paragraph()

# ─── 5. SUSPICIOUS PLAYER DETECTION ─────────────────────────────────────────
doc.add_page_break()
add_heading(doc, '5. Suspicious Player Detection Logic', level=1, color=(31, 56, 100))

add_paragraph(doc,
    'Detection uses a two-layer approach: explicit rule-based thresholds (instant, '
    'explainable, debuggable live in a demo) and historical statistical Z-score analysis '
    '(catches slow-building anomalies). No ML model training data is required.')

doc.add_paragraph()
add_paragraph(doc, '5.1  Per-Match Rule Checks (Applied to Every Row)', bold=True)

cheat_rules = [
    ('Rule 1', 'Score Rate Spike',
     'score_per_second > 50',
     'High',
     'Normal gameplay: ~7–10 score/sec. P003 example: 99,000 ÷ 60s = 1,650 score/sec — clearly impossible.'),
    ('Rule 2', 'Impossible K/D',
     'kills ≥ 20 AND deaths == 0',
     'High',
     'Zero deaths with high kills is statistically near-impossible in a real match.'),
    ('Rule 3', 'Short Match + High Score',
     'duration < 120s AND score > 5,000',
     'Medium',
     'A very short match cannot legitimately produce a high score under normal game mechanics.'),
    ('Rule 4', 'Extreme Kill Count',
     'kills > 50 in one match',
     'Medium',
     'Based on realistic game design caps; adjust per game type meta.'),
    ('Rule 5', 'Ping Anomaly',
     'ping < 5ms OR ping > 500ms',
     'Low',
     'Data-quality flag, not necessarily cheating. Indicates network anomaly or data corruption.'),
]

tbl5 = doc.add_table(rows=1, cols=5)
tbl5.style = 'Table Grid'
for i, txt in enumerate(['Rule', 'Name', 'Condition', 'Severity', 'Rationale']):
    cell = tbl5.rows[0].cells[i]
    cell.text = txt
    set_cell_bg(cell, '7030A0')
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255,255,255)
            run.font.size = Pt(9)

sev_colors = {'High': 'FFCCCC', 'Medium': 'FFE5B4', 'Low': 'FFFACD'}
for rule, name, cond, sev, rat in cheat_rules:
    row = tbl5.add_row()
    row.cells[0].text = rule
    row.cells[1].text = name
    row.cells[2].text = cond
    row.cells[3].text = sev
    row.cells[4].text = rat
    set_cell_bg(row.cells[3], sev_colors.get(sev, 'FFFFFF'))
    for c in row.cells:
        c.paragraphs[0].runs[0].font.size = Pt(9)
    row.cells[0].paragraphs[0].runs[0].font.bold = True

doc.add_paragraph()
add_paragraph(doc, '5.2  Historical Z-Score Outlier Detection', bold=True)
add_paragraph(doc,
    'For players with 3 or more recorded matches, we compute per-player statistics and '
    'flag any match where the player\'s own score or kill count deviates sharply '
    'from their personal history:')
add_bullet(doc, 'z_score = (match_value − player_mean) ÷ player_std')
add_bullet(doc, 'If |z_score| > 3.0 for score OR kills → flag as "statistical outlier" (Low severity)')
add_bullet(doc, 'This catches players who are normally average but had one suspiciously extraordinary match.')
add_bullet(doc, 'No cross-player data or training required — purely intra-player statistics.')

doc.add_paragraph()
add_paragraph(doc, '5.3  Severity Escalation Rules', bold=True)
add_bullet(doc, 'High:   Rule 1 or Rule 2 triggered (near-impossible stats)')
add_bullet(doc, 'Medium: Rule 3 or Rule 4 triggered (suspicious but not impossible)')
add_bullet(doc, 'Low:    Only Z-score deviation or Rule 5 (anomalous, worth monitoring)')
add_bullet(doc, 'Escalation Override: If ≥ 5 matches played AND > 50% trigger gameplay rules → auto-escalate to High')

doc.add_paragraph()
add_paragraph(doc, '5.4  Worked Example — Player P003', bold=True)

example_rows = [
    ('player_id', 'match_id', 'score', 'kills', 'deaths', 'duration', 'score/sec', 'Rules Triggered'),
    ('P003', 'M002', '99,000', '250', '0', '60s', '1,650', 'Rule 1 ✓  Rule 2 ✓  Rule 3 ✓  Rule 4 ✓'),
    ('P001', 'M001', '3,200', '18', '4', '420s', '7.6', 'None — Clean'),
]
tbl6 = doc.add_table(rows=0, cols=8)
tbl6.style = 'Table Grid'
for r_idx, row_data in enumerate(example_rows):
    row = tbl6.add_row()
    for c_idx, val in enumerate(row_data):
        row.cells[c_idx].text = val
        if r_idx == 0:
            set_cell_bg(row.cells[c_idx], '4472C4')
            for para in row.cells[c_idx].paragraphs:
                for run in para.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255,255,255)
                    run.font.size = Pt(9)
        elif 'P003' in val or '✓' in val:
            set_cell_bg(row.cells[c_idx], 'FFCCCC')
            row.cells[c_idx].paragraphs[0].runs[0].font.size = Pt(9)
        else:
            set_cell_bg(row.cells[c_idx], 'E2EFDA')
            row.cells[c_idx].paragraphs[0].runs[0].font.size = Pt(9)

doc.add_paragraph()

# ─── 6. MATCHMAKING LOGIC ────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, '6. Matchmaking Logic', level=1, color=(31, 56, 100))

add_paragraph(doc,
    'Players are grouped into balanced lobbies using a three-pass bucketing algorithm. '
    'The system prioritises fairness (similar skill), latency quality (similar ping), '
    'and safety (cheater isolation).')

doc.add_paragraph()
add_paragraph(doc, '6.1  Skill Score Calculation', bold=True)
add_paragraph(doc,
    'A simple weighted skill metric is computed per player from their clean match history:')
p_skill = doc.add_paragraph()
p_skill.paragraph_format.left_indent = Inches(0.5)
run = p_skill.add_run(
    'skill_score = (avg_score × 0.5) + (avg_kills × 3.0) − (avg_deaths × 1.0)')
run.font.bold = True
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(31, 56, 100)

add_paragraph(doc,
    'This is not Elo — it is a simple ordinal skill estimate sufficient for percentile '
    'bucketing at prototype scale. Weights can be tuned per game genre.')

doc.add_paragraph()
add_paragraph(doc, '6.2  Three-Pass Bucketing Algorithm', bold=True)

passes = [
    ('Pass 1 — Region',
     'Group players by primary_region (India, SEA, Europe…). '
     'Same-region players share similar network baseline ping, producing fair latency conditions.'),
    ('Pass 2 — Ping Band',
     'Subdivide each region group into three ping tiers:\n'
     '  • Low Ping   : ping ≤ 80ms\n'
     '  • Medium Ping: 81ms – 150ms\n'
     '  • High Ping  : > 150ms\n'
     'Hard rule: Low Ping players are NEVER matched with High Ping players.'),
    ('Pass 3 — Skill Tier',
     'Within each Region + Ping Band bucket:\n'
     '  • If bucket has ≥ 6 players → split by percentile: Beginner (bottom 33%), Intermediate (middle 33%), Advanced (top 33%).\n'
     '  • If bucket has < 6 players → keep as "Mixed" pool (small pool exception).\n'
     '  • Minimum lobby size of 4 players for "Ready" status; smaller lobbies are marked "Waiting".'),
]

for title_p, desc in passes:
    add_paragraph(doc, title_p, bold=True, indent=True)
    add_paragraph(doc, desc, indent=True)

doc.add_paragraph()
add_paragraph(doc, '6.3  Handling Flagged Players in Matchmaking', bold=True)

sev_handle = [
    ('High Severity',   'Excluded completely from all match lobbies. Cannot be matched with clean players.'),
    ('Medium Severity', 'Placed in an isolated "Under_Review" lobby. Kept separate pending manual investigation.'),
    ('Low Severity',    'Matched normally within standard lobbies but tagged in output for monitoring.'),
]
tbl7 = doc.add_table(rows=1, cols=2)
tbl7.style = 'Table Grid'
for i, txt in enumerate(['Severity Level', 'Matchmaking Treatment']):
    cell = tbl7.rows[0].cells[i]
    cell.text = txt
    set_cell_bg(cell, '833C00')
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255,255,255)
            run.font.size = Pt(10)
for sev, treatment in sev_handle:
    row = tbl7.add_row()
    row.cells[0].text = sev
    row.cells[1].text = treatment
    c0 = {'High Severity': 'FFCCCC', 'Medium Severity': 'FFE5B4', 'Low Severity': 'FFFACD'}
    set_cell_bg(row.cells[0], c0.get(sev, 'FFFFFF'))
    for c in row.cells:
        c.paragraphs[0].runs[0].font.size = Pt(10)
    row.cells[0].paragraphs[0].runs[0].font.bold = True

doc.add_paragraph()

# ─── 7. ARCHITECTURE EXPLANATION ─────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, '7. Architecture & Technical Choices', level=1, color=(31, 56, 100))

add_paragraph(doc, '7.1  System Architecture Overview', bold=True)
add_paragraph(doc,
    'The system uses a modular, pipeline-based architecture where each functional concern '
    '(loading, detection, ranking, matchmaking) is isolated into its own Python module. '
    'This makes each component independently testable, replaceable, and explainable.')

doc.add_paragraph()
arch_flow = [
    ('Input Layer',      'data/players.csv',              'Raw match telemetry — single CSV input source'),
    ('Loader Module',    'src/loader.py',                  'Validates, cleans, normalises, deduplicates, computes derived columns'),
    ('Detection Module', 'src/suspicious.py',              'Applies 5 cheat rules + Z-score history per player'),
    ('Leaderboard',      'src/leaderboard.py',             'Aggregates stats, sorts with tie-breakers, splits global & regional'),
    ('Matchmaking',      'src/matchmaking.py',             'Computes skill scores, 3-pass bucketing, isolates cheaters'),
    ('CLI Runner',       'main.py',                        'Orchestrates full pipeline, writes 4 CSV reports to output/'),
    ('API Layer',        'api/routes.py + api/app.py',     'FastAPI REST endpoints (POST /submit-score, GET /leaderboard, etc.)'),
    ('Web Dashboard',    'static/ (HTML + CSS + JS)',      'Premium glassmorphic dark-mode single-page application'),
]

tbl8 = doc.add_table(rows=1, cols=3)
tbl8.style = 'Table Grid'
for i, txt in enumerate(['Layer', 'File / Component', 'Responsibility']):
    cell = tbl8.rows[0].cells[i]
    cell.text = txt
    set_cell_bg(cell, '1F3864')
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255,255,255)
            run.font.size = Pt(10)
for layer, comp, resp in arch_flow:
    row = tbl8.add_row()
    row.cells[0].text = layer
    row.cells[1].text = comp
    row.cells[2].text = resp
    for c in row.cells:
        c.paragraphs[0].runs[0].font.size = Pt(10)
    row.cells[0].paragraphs[0].runs[0].font.bold = True
    set_cell_bg(row.cells[0], 'DCE6F1')

doc.add_paragraph()
add_paragraph(doc, '7.2  Technology Choices', bold=True)

tech = [
    ('Python + Pandas', 'Data wrangling, grouping, filtering — ideal for tabular match data; zero-boilerplate aggregations.'),
    ('NumPy',           'Z-score arithmetic and vectorised statistical calculations.'),
    ('FastAPI',         'High-performance async REST framework; auto-generates OpenAPI docs; Pydantic validation on every endpoint.'),
    ('Uvicorn',         'ASGI server to run FastAPI — handles async concurrency efficiently.'),
    ('HTML/CSS/JS',     'Vanilla frontend (no framework overhead); premium glassmorphic dark-mode design with Google Fonts.'),
    ('pytest',          'Unit test coverage for all core logic — 5 tests, all passing.'),
    ('python-docx',     'Generates this document programmatically — ensuring reproducibility.'),
]

tbl9 = doc.add_table(rows=1, cols=2)
tbl9.style = 'Table Grid'
for i, txt in enumerate(['Technology', 'Why Chosen']):
    cell = tbl9.rows[0].cells[i]
    cell.text = txt
    set_cell_bg(cell, '203864')
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255,255,255)
            run.font.size = Pt(10)
for tech_name, why in tech:
    row = tbl9.add_row()
    row.cells[0].text = tech_name
    row.cells[1].text = why
    for c in row.cells:
        c.paragraphs[0].runs[0].font.size = Pt(10)
    row.cells[0].paragraphs[0].runs[0].font.bold = True

doc.add_paragraph()

# ─── 8. SCALING PLAN ─────────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, '8. Scaling Plan — 2,000 to 200,000+ Players', level=1, color=(31, 56, 100))

add_paragraph(doc,
    'The prototype uses a synchronous in-memory pipeline. For production scale, we transition '
    'to a fully decoupled, event-driven architecture:')

doc.add_paragraph()
scale_points = [
    ('Database — PostgreSQL',
     'Replace Pandas in-memory DataFrame with a PostgreSQL cluster. Tables are horizontally '
     'partitioned by region and created_at date — this keeps each partition small, ensuring '
     'index lookups and range queries remain fast even at 200M rows.'),
    ('Leaderboard Cache — Redis Sorted Sets',
     'Direct SQL ORDER BY on 200,000 players is expensive. Redis Sorted Sets (ZADD / ZREVRANGE) '
     'update and serve player ranks in O(log N) time with p99 read latency < 5ms. A background '
     'job refreshes the sorted set from PostgreSQL every 30 seconds.'),
    ('Async Ingestion — Kafka / AWS SQS',
     'Match score events are published to a message queue immediately on receipt by the API '
     '(< 10ms response). Async consumer workers batch-write to PostgreSQL, decoupling write '
     'throughput from API response latency. Handles 200,000+ match submissions per day.'),
    ('API Horizontal Scaling',
     'FastAPI pods are stateless and containerised (Docker). A load balancer auto-scales '
     'pod count based on request rate. Kubernetes HPA (Horizontal Pod Autoscaler) spins up '
     'new pods within 30 seconds of a traffic spike.'),
    ('Real-Time Anti-Cheat — Apache Flink',
     'Rule-based checks (Rules 1–5) run as a streaming job directly on the Kafka topic, '
     'firing alerts in real time as match events arrive — no batch delay.'),
    ('Batch Z-Score Detection — Apache Spark',
     'Historical Z-score analysis runs as a scheduled Spark job every 10 minutes over the '
     'PostgreSQL match history, updating player-level anomaly flags without real-time pressure.'),
    ('Monitoring — Prometheus + Grafana',
     'Prometheus scrapes API latency, queue consumer lag, error rates, and DB connection counts. '
     'Grafana dashboards give ops teams live visibility into system health with alerting thresholds.'),
]

for component, detail in scale_points:
    add_paragraph(doc, component, bold=True)
    add_paragraph(doc, detail, indent=True)
    doc.add_paragraph()

# Quick table summary
add_paragraph(doc, 'Scaling Capability Summary', bold=True)
scale_summary = [
    ('2,000 players',   'Current prototype — in-memory Pandas, single FastAPI instance'),
    ('20,000 players',  'PostgreSQL + Redis cache + multiple API pods behind load balancer'),
    ('200,000 players', 'Kafka ingestion queue + Flink stream processing + Spark batch analytics + Redis sorted sets'),
    ('2,000,000+ players', 'Multi-region deployment, read replicas, CDN-cached leaderboard snapshots'),
]
tbl10 = doc.add_table(rows=1, cols=2)
tbl10.style = 'Table Grid'
for i, txt in enumerate(['Scale', 'Architecture Tier']):
    cell = tbl10.rows[0].cells[i]
    cell.text = txt
    set_cell_bg(cell, '1F3864')
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255,255,255)
            run.font.size = Pt(10)
for scale, arch in scale_summary:
    row = tbl10.add_row()
    row.cells[0].text = scale
    row.cells[1].text = arch
    for c in row.cells:
        c.paragraphs[0].runs[0].font.size = Pt(10)
    row.cells[0].paragraphs[0].runs[0].font.bold = True

doc.add_paragraph()

# ─── 9. LIMITATIONS ──────────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, '9. Limitations', level=1, color=(31, 56, 100))

limitations = [
    ('Static Thresholds',
     'Anti-cheat rule limits (score/sec > 50, kills > 50, etc.) are manually tuned to the '
     'provided sample data. In production, these should be dynamically calibrated per game '
     'version, weapon balance patch, and event type using historical distribution analysis.'),
    ('No API Authentication',
     'REST endpoints currently accept any request without token validation or rate limiting. '
     'Production systems must implement HMAC-signed requests from verified game clients and '
     'OAuth2 for operations team dashboards.'),
    ('Static Matchmaking (No Live Queue)',
     'The matchmaking algorithm is a one-shot bucketing pass on current data. A production '
     'system needs a live lobby queue (e.g., Open Match) that dynamically relaxes skill '
     'constraints if a player has been waiting too long.'),
    ('In-Memory State — No Persistence',
     'The FastAPI server holds the dataset in a Pandas DataFrame. A server restart wipes '
     'all submitted scores. Production requires a persistent database with transaction logging.'),
    ('Z-Score Requires Minimum History',
     'Statistical outlier detection only activates after a player has ≥ 3 matches. New players '
     'are not covered by the historical layer — only real-time rule checks apply to them.'),
    ('Region Normalisation',
     '"SEA" is displayed as "Sea" after title-case normalisation. A canonical region registry '
     'with explicit mappings (e.g., "SEA" → "Southeast Asia") should be used in production.'),
    ('No Season Reset Mechanism',
     'The leaderboard is cumulative and does not support season-based resets. This would '
     'require adding a season_id column and filtering queries by active season window.'),
]

for i, (title_l, desc) in enumerate(limitations, 1):
    add_paragraph(doc, f'{i}. {title_l}', bold=True)
    add_paragraph(doc, desc, indent=True)
    doc.add_paragraph()

# ─── 10. DEMO SCRIPT ─────────────────────────────────────────────────────────
add_heading(doc, '10. 5-Minute Demo Walkthrough', level=1, color=(31, 56, 100))

demo_steps = [
    ('0:00 – 0:30', 'Show Input CSV',
     'Open data/players.csv. Point to P003: score=99,000, kills=250, deaths=0, duration=60s. '
     'Ask audience: "Does this look realistic?"'),
    ('0:30 – 1:00', 'Run Pipeline',
     'Execute: python main.py\nShow console output — 18 rows loaded, 1 player flagged, 7 lobbies formed.'),
    ('1:00 – 2:00', 'Show Anti-Cheat Output',
     'Open output/suspicious_players.csv. Explain Rule 1 caught P003 at 1,650 score/sec. '
     'Explain severity = High → excluded from leaderboard ranks.'),
    ('2:00 – 3:00', 'Show Leaderboard',
     'Open output/leaderboard_global.csv. P001 is Rank 1 (3 matches, 9,700 total score). '
     'P003 has score 99,000 but rank = blank. Explain: "P003\'s score doesn\'t move anyone down."'),
    ('3:00 – 4:00', 'Show Matchmaking',
     'Open output/matchmaking_groups.csv. Explain India_LowPing groups vs SEA_MedPing_Mixed '
     'vs Europe_HighPing_Mixed. Show P003 is absent from all groups.'),
    ('4:00 – 5:00', 'Scaling & Trade-offs',
     'Briefly: Redis sorted sets for leaderboard reads in < 5ms, Kafka queue for ingestion, '
     'stateless API pods behind load balancer. Trade-offs: thresholds need tuning, no live queue yet.'),
]

tbl11 = doc.add_table(rows=1, cols=3)
tbl11.style = 'Table Grid'
for i, txt in enumerate(['Time', 'Step', 'What to Show']):
    cell = tbl11.rows[0].cells[i]
    cell.text = txt
    set_cell_bg(cell, '1F3864')
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255,255,255)
            run.font.size = Pt(10)
for timing, step, show in demo_steps:
    row = tbl11.add_row()
    row.cells[0].text = timing
    row.cells[1].text = step
    row.cells[2].text = show
    row.cells[0].paragraphs[0].runs[0].font.bold = True
    for c in row.cells:
        c.paragraphs[0].runs[0].font.size = Pt(10)

doc.add_paragraph()

# ─── FILE LIST ────────────────────────────────────────────────────────────────
add_heading(doc, '11. Project File Structure', level=1, color=(31, 56, 100))

files_block = (
    'Game/\n'
    '├── data/\n'
    '│   └── players.csv                 ← Sample match input data\n'
    '├── src/\n'
    '│   ├── loader.py                   ← Data ingestion & validation\n'
    '│   ├── suspicious.py               ← Anti-cheat detection engine\n'
    '│   ├── leaderboard.py              ← Ranking & aggregation\n'
    '│   ├── matchmaking.py              ← 3-pass lobby bucketing\n'
    '│   └── utils.py                    ← Shared helpers\n'
    '├── api/\n'
    '│   ├── app.py                      ← FastAPI application entry point\n'
    '│   └── routes.py                   ← REST API endpoints\n'
    '├── static/\n'
    '│   ├── index.html                  ← Web dashboard HTML\n'
    '│   ├── styles.css                  ← Glassmorphic dark-mode styling\n'
    '│   └── dashboard.js                ← Dashboard controller & API client\n'
    '├── tests/\n'
    '│   └── test_logic.py               ← pytest unit tests (5 tests, all pass)\n'
    '├── output/\n'
    '│   ├── leaderboard_global.csv\n'
    '│   ├── leaderboard_region.csv\n'
    '│   ├── suspicious_players.csv\n'
    '│   └── matchmaking_groups.csv\n'
    '├── main.py                         ← CLI pipeline runner\n'
    '├── design_doc.md                   ← Markdown architecture document\n'
    '└── requirements.txt'
)

p_code = doc.add_paragraph()
p_code.paragraph_format.left_indent = Inches(0.3)
run_code = p_code.add_run(files_block)
run_code.font.name = 'Courier New'
run_code.font.size = Pt(9)

doc.add_paragraph()

# ─── FOOTER / SIGNATURE ───────────────────────────────────────────────────────
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer_p.add_run(
    '— End of Technical Document — \n'
    'Game Ops System | Hybrid Rule-Based Python Approach\n'
    'pytest: 5/5 Tests Passed | API: FastAPI + Uvicorn | Dashboard: Glassmorphic SPA'
)
footer_run.italic = True
footer_run.font.size = Pt(10)
footer_run.font.color.rgb = RGBColor(120, 120, 120)

# ─── SAVE ─────────────────────────────────────────────────────────────────────
out_path = os.path.join('output', 'GameOps_Technical_Document.docx')
os.makedirs('output', exist_ok=True)
doc.save(out_path)
print(f"[OK]  Document saved to: {out_path}")
