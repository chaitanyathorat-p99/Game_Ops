from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── helpers ──────────────────────────────────────────────────────────────────

def cell_bg(cell, hex_col):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_col)
    tcPr.append(shd)

def h(doc, text, lvl=2, rgb=(31,56,100)):
    hd = doc.add_heading(text, level=lvl)
    hd.paragraph_format.space_before = Pt(4)
    hd.paragraph_format.space_after  = Pt(1)
    r = hd.runs[0]
    r.font.color.rgb = RGBColor(*rgb)
    r.font.bold = True
    r.font.size = Pt(9 if lvl >= 2 else 11)
    return hd

def p(doc, text, bold=False, size=8, indent=0, space_after=2):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(space_after)
    para.paragraph_format.left_indent  = Inches(indent)
    r = para.add_run(text)
    r.bold      = bold
    r.font.size = Pt(size)
    return para

def bullet(doc, text, size=8, indent=0.15):
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(1)
    para.paragraph_format.left_indent  = Inches(indent)
    r = para.add_run(text)
    r.font.size = Pt(size)
    return para

def mini_table(doc, headers, rows, hdr_rgb='1F3864', col_widths=None):
    ncols = len(headers)
    tbl   = doc.add_table(rows=1, cols=ncols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.paragraph_format_space_after = Pt(2)

    # header row
    for i, hdr in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = hdr
        cell_bg(cell, hdr_rgb)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold  = True
                run.font.size  = Pt(7.5)
                run.font.color.rgb = RGBColor(255, 255, 255)

    # data rows
    alt = ['F2F7FF', 'FFFFFF']
    for ri, row_data in enumerate(rows):
        row = tbl.add_row()
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = val
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell_bg(cell, alt[ri % 2])
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(7.5)

    # optional column widths
    if col_widths:
        for row in tbl.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Inches(w)
    return tbl

def spacer(doc, pts=3):
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after  = Pt(pts)

# ── document ──────────────────────────────────────────────────────────────────

doc = Document()

# Tight margins for 2-page fit
for sec in doc.sections:
    sec.top_margin    = Cm(1.4)
    sec.bottom_margin = Cm(1.4)
    sec.left_margin   = Cm(1.8)
    sec.right_margin  = Cm(1.8)
    sec.page_height   = Cm(29.7)   # A4
    sec.page_width    = Cm(21.0)

# default style
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(8)

# ── TITLE ─────────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(1)
r1 = title.add_run('GAME OPS SYSTEM  —  Technical Design Document')
r1.bold = True
r1.font.size = Pt(13)
r1.font.color.rgb = RGBColor(31, 56, 100)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(4)
r2 = sub.add_run('Multiplayer Action Game | Weekend Event Operations | Rule-Based Python + FastAPI + Web Dashboard')
r2.italic = True
r2.font.size = Pt(8)
r2.font.color.rgb = RGBColor(90, 90, 90)

# ── 1. PROBLEM UNDERSTANDING ──────────────────────────────────────────────────
h(doc, '1.  Problem Understanding')
p(doc,
  'During a live weekend event, thousands of players submit match results simultaneously. '
  'The ops team needs a system that: (a) builds a fair leaderboard while isolating '
  'fraudulent scores, (b) detects cheaters without requiring ML training data, '
  '(c) groups players into low-latency, skill-balanced lobbies, and (d) explains how '
  'the system scales from 2,000 to 200,000+ concurrent players.',
  size=8)

spacer(doc, 3)

# ── 2. ASSUMPTIONS ────────────────────────────────────────────────────────────
h(doc, '2.  Assumptions')
mini_table(doc,
    ['Area', 'Assumption'],
    [
        ('Data source',    'CSV / REST POST — no real-time stream at prototype level'),
        ('Thresholds',     'Manually tuned to sample data; configurable per game meta'),
        ('Fair leaderboard','Flagged players shown as "under_review" — do NOT displace clean ranks'),
        ('Matchmaking',    'Static bucketing sufficient at prototype; live queue is a prod enhancement'),
        ('Persistence',    'In-memory Pandas DataFrame; prod uses PostgreSQL + Redis'),
        ('Auth / rate-limit','Not implemented at prototype; required in production'),
    ],
    hdr_rgb='1F3864',
    col_widths=[1.4, 4.9]
)

spacer(doc, 3)

# ── 3. DATA STRUCTURE & DATABASE DESIGN ──────────────────────────────────────
h(doc, '3.  Data Structure & Database Design')
p(doc, 'Two logical tables (implemented as Pandas DataFrames; mapped to PostgreSQL in production):', size=8)

p(doc, 'matches  (raw telemetry)', bold=True, size=8, space_after=1)
mini_table(doc,
    ['Field', 'Type', 'Note'],
    [
        ('player_id / match_id',       'VARCHAR',  'Composite primary key — deduplicated on ingest'),
        ('region / device',            'VARCHAR',  'Normalised to title-case on load'),
        ('ping / score / kills / deaths / duration', 'INT / FLOAT', 'Validated: duration > 0, all fields >= 0'),
        ('score_per_second',           'FLOAT',    'DERIVED — score ÷ duration'),
        ('kd_ratio',                   'FLOAT',    'DERIVED — kills ÷ max(deaths, 1)'),
    ],
    hdr_rgb='2E75B6',
    col_widths=[2.3, 1.1, 2.9]
)
spacer(doc, 2)
p(doc, 'players  (aggregated, updated per batch)', bold=True, size=8, space_after=1)
mini_table(doc,
    ['Field', 'Description'],
    [
        ('total_score / kills / deaths / matches_played', 'Running aggregates across all clean matches'),
        ('skill_score',    'Weighted metric: avg_score×0.5 + avg_kills×3.0 − avg_deaths×1.0'),
        ('cheat_status / cheat_severity / flag_reasons',  '"clean" or "under_review" with severity level and triggered rules'),
    ],
    hdr_rgb='2E75B6',
    col_widths=[2.8, 3.5]
)

spacer(doc, 3)

# ── 4. LEADERBOARD LOGIC ──────────────────────────────────────────────────────
h(doc, '4.  Leaderboard Logic')
p(doc,
  'Player stats are aggregated per match history, then sorted using a cascading tie-breaker:',
  size=8)
mini_table(doc,
    ['Priority', 'Column', 'Direction', 'Rationale'],
    [
        ('1st', 'total_score',    'DESC', 'Higher cumulative score = better performance'),
        ('2nd', 'total_deaths',   'ASC',  'Fewer deaths wins on equal score'),
        ('3rd', 'total_kills',    'DESC', 'More kills wins if deaths also equal'),
        ('4th', 'matches_played', 'DESC', 'Consistency rewarded if all else equal'),
    ],
    hdr_rgb='375623',
    col_widths=[0.7, 1.4, 1.0, 3.2]
)
spacer(doc, 1)
p(doc,
  'Flagged players appear at the bottom with rank = None and status = "under_review". '
  'Their scores do NOT shift any clean player\'s rank number. A separate regional '
  'leaderboard applies the same logic within each primary_region group.',
  size=8)

spacer(doc, 3)

# ── 5. SUSPICIOUS PLAYER DETECTION ───────────────────────────────────────────
h(doc, '5.  Suspicious Player Detection Logic')
p(doc, 'Two-layer approach — no ML training data needed:', size=8)
p(doc, 'Layer 1 — Per-match rule checks:', bold=True, size=8, space_after=1)
mini_table(doc,
    ['Rule', 'Condition', 'Severity', 'Example'],
    [
        ('R1 Score spike',  'score_per_second > 50',           'High',   'P003: 99000/60s = 1650/s'),
        ('R2 K/D zero-death','kills >= 20 AND deaths == 0',    'High',   'P003: 250 kills, 0 deaths'),
        ('R3 Short+high',   'duration < 120s AND score > 5000','Medium', 'Score 99000 in 60s'),
        ('R4 Extreme kills','kills > 50 in one match',         'Medium', '250 kills in one match'),
        ('R5 Ping anomaly', 'ping < 5ms OR ping > 500ms',      'Low',    'Network/data quality flag'),
    ],
    hdr_rgb='7030A0',
    col_widths=[1.3, 2.0, 0.8, 2.2]
)
spacer(doc, 1)
p(doc, 'Layer 2 — Historical Z-score (>=3 matches needed):', bold=True, size=8, space_after=1)
p(doc,
  'z = (match_value − player_mean) ÷ player_std.  |z| > 3.0 for score OR kills → "Low" severity flag. '
  'Catches players who are normally average but had one extraordinary match. '
  'Escalation: if >= 5 matches AND > 50% trigger gameplay rules → severity overridden to High.',
  size=8)

spacer(doc, 3)

# ── 6. MATCHMAKING LOGIC ──────────────────────────────────────────────────────
h(doc, '6.  Matchmaking Logic')
p(doc, '3-pass bucketing algorithm:', bold=True, size=8, space_after=1)
mini_table(doc,
    ['Pass', 'Criterion', 'Rule'],
    [
        ('Pass 1', 'Region',     'Group by primary region — minimises baseline ping variance'),
        ('Pass 2', 'Ping band',  'Low ≤80ms | Med 81-150ms | High >150ms — Low/High never mixed'),
        ('Pass 3', 'Skill tier', '>=6 players: Beginner/Intermediate/Advanced (33rd percentile splits); <6: Mixed pool'),
    ],
    hdr_rgb='833C00',
    col_widths=[0.6, 1.1, 4.6]
)
spacer(doc, 1)
p(doc,
  'Cheater handling — High severity: excluded entirely | Medium: isolated "Under_Review" lobby | '
  'Low: matched normally but tagged. Lobbies with < 4 players are marked "Waiting".',
  size=8)

spacer(doc, 3)

# ── 7. ARCHITECTURE ───────────────────────────────────────────────────────────
h(doc, '7.  Architecture Explanation')
p(doc, 'Modular pipeline — each concern is an isolated, independently testable Python module:', size=8)
mini_table(doc,
    ['Layer', 'Component', 'Responsibility'],
    [
        ('Input',      'data/players.csv',        'Raw match telemetry — single CSV source'),
        ('Loader',     'src/loader.py',            'Validate, clean, deduplicate, compute derived columns'),
        ('Detection',  'src/suspicious.py',        'Rule checks (R1–R5) + Z-score history per player'),
        ('Leaderboard','src/leaderboard.py',       'Aggregate, sort with tie-breakers, split global/regional'),
        ('Matchmaking','src/matchmaking.py',       'Skill score, 3-pass bucketing, isolate cheaters'),
        ('CLI',        'main.py',                  'Orchestrate full pipeline → 4 CSV reports in output/'),
        ('API',        'api/ (FastAPI + uvicorn)', 'POST /submit-score | GET /leaderboard | /flagged-players | /matchmaking | /download'),
        ('Dashboard',  'static/ (HTML+CSS+JS)',    'Premium glassmorphic SPA — real-time metrics, upload, tables, lobbies'),
    ],
    hdr_rgb='1F3864',
    col_widths=[1.1, 1.8, 3.4]
)

spacer(doc, 3)

# ── 8. SCALING PLAN ───────────────────────────────────────────────────────────
h(doc, '8.  Scaling Plan')
mini_table(doc,
    ['Scale', 'Component', 'Approach'],
    [
        ('2K players',    'Current prototype',    'In-memory Pandas + single FastAPI pod'),
        ('20K players',   'Database + Cache',     'PostgreSQL (partitioned by region/date) + Redis Sorted Set for leaderboard reads < 5ms'),
        ('200K players',  'Async ingestion',      'Kafka / AWS SQS message queue — API publishes events; async workers write to DB, decoupling write pressure'),
        ('200K players',  'Real-time anti-cheat', 'Apache Flink stream job on Kafka topic — fires rule alerts in real time'),
        ('200K players',  'Batch Z-score',        'Apache Spark scheduled job every 10 min over PostgreSQL match history'),
        ('Any scale',     'API pods',             'Stateless FastAPI containers; Kubernetes HPA auto-scales on request rate'),
        ('Any scale',     'Monitoring',           'Prometheus scrapes latency/queue lag/errors; Grafana dashboards with alert thresholds'),
    ],
    hdr_rgb='154360',
    col_widths=[1.0, 1.6, 3.7]
)

spacer(doc, 3)

# ── 9. LIMITATIONS ────────────────────────────────────────────────────────────
h(doc, '9.  Limitations')
mini_table(doc,
    ['Limitation', 'Detail'],
    [
        ('Static thresholds',       'Rule limits are manually tuned; prod needs auto-calibration per game version/patch'),
        ('No authentication',       'API endpoints have no token validation or rate limiting — required in production'),
        ('Static matchmaking',      'One-shot bucketing; no live queue with wait-time relaxation'),
        ('No persistence on restart','In-memory DataFrame lost on server restart; prod uses DB with transaction log'),
        ('Z-score needs history',   'New players (< 3 matches) not covered by statistical layer — rule checks only'),
        ('Region normalisation',    '"SEA" becomes "Sea" after title-case; prod needs canonical region registry'),
        ('No season reset',         'Leaderboard is cumulative; season_id and date-window filtering needed for resets'),
    ],
    hdr_rgb='7B241C',
    col_widths=[1.8, 4.5]
)

spacer(doc, 2)

# ── footer ─────────────────────────────────────────────────────────────────────
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
foot.paragraph_format.space_before = Pt(4)
fr = foot.add_run(
    'Game Ops System  |  Hybrid Rule-Based Python Approach  |  '
    'pytest: 5/5 Tests Passed  |  Stack: Python + FastAPI + Pandas + Vanilla JS')
fr.italic = True
fr.font.size = Pt(7)
fr.font.color.rgb = RGBColor(140, 140, 140)

# ── save ─────────────────────────────────────────────────────────────────────
out = os.path.join('output', 'GameOps_2Page_Document.docx')
os.makedirs('output', exist_ok=True)
doc.save(out)
print(f"[OK] Document saved -> {out}")
